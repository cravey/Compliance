import sys
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER

def format_text_with_lists(doc, text):
    lines = text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith(('1)', '2)', '3)', 'a.', 'b.', 'c.', 'i.', 'ii.', 'iii.')):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped_line)
            run.font.size = Pt(12)
            if stripped_line[0].isdigit():  # e.g., '1)', '2)', etc.
                paragraph.paragraph_format.left_indent = Inches(1)
            elif stripped_line[0].isalpha():  # e.g., 'a.', 'b.', etc.
                paragraph.paragraph_format.left_indent = Inches(0.5)
        else:
            doc.add_paragraph(line)

def generate_document(control_id, control, control_text, discussion, related_controls, actions, milestones, resources, system_name, owner):
    doc = Document()

    # Upper left: Control ID and Upper right: System Name on the same line
    header_para = doc.add_paragraph()
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    run = header_para.add_run(control_id)
    run.font.size = Pt(24)
    header_para.add_run('\t')
    run = header_para.add_run(system_name)
    run.font.size = Pt(24)

    # Below System Name: Owner
    owner_para = doc.add_paragraph()
    run = owner_para.add_run(f"Owner: {owner}")
    run.font.size = Pt(12)
    owner_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Centered: Control
    control_para = doc.add_paragraph()
    run = control_para.add_run(control)
    run.font.size = Pt(18)
    control_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Control Text
    doc.add_heading('Control Text:', level=2)
    format_text_with_lists(doc, control_text)

    # Discussion
    doc.add_heading('Discussion:', level=2)
    format_text_with_lists(doc, discussion)

    # Related Controls
    doc.add_heading('Related Controls:', level=2)
    if related_controls.strip() == "":
        related_controls = "None"
    format_text_with_lists(doc, related_controls)

    # Horizontal line
    doc.add_paragraph().add_run().add_break()
    doc.add_paragraph().add_run('__________________________').bold = True

    # List of actions
    doc.add_heading('List of Actions', level=2)
    for action in actions:
        doc.add_paragraph(action)

    # List of milestones with estimated completion dates
    doc.add_heading('List of Milestones with Estimated Completion Dates', level=2)
    for milestone in milestones:
        doc.add_paragraph(milestone)

    # List of required resources
    doc.add_heading('List of Required Resources', level=2)
    for resource in resources:
        doc.add_paragraph(resource)

    # Save document
    doc.save(f"{control_id}.docx")

def main():
    control_id = "AC-2"
    control = "Account Management"
    control_text = ("The organization: a. Identifies and selects the following types of information system accounts to support organizational missions/business functions: "
                    "individual, group, system, application, guest/anonymous, and temporary; b. Assigns account managers for information system accounts; c. Establishes conditions for group "
                    "membership; d. Specifies authorized users of the information system, group and role membership, and access authorizations (i.e., privileges) and other attributes (as required) "
                    "for each account; e. Requires approvals by [Assignment: organization-defined personnel or roles] for requests to create information system accounts; f. Creates, enables, modifies, "
                    "disables, and removes information system accounts in accordance with [Assignment: organization-defined procedures or conditions]; g. Monitors the use of information system accounts; "
                    "h. Notifies account managers: 1. When accounts are no longer required; 2. When users are terminated or transferred; and 3. When individual information system usage or need-to-know "
                    "changes; i. Authorizes access to the information system based on: 1. A valid access authorization; 2. Intended system usage; and 3. Other attributes as required by the organization "
                    "or associated missions/business functions; j. Reviews accounts for compliance with account management requirements [Assignment: organization-defined frequency]; and k. Establishes a process "
                    "for reissuing shared/group account credentials (if deployed) when individuals are removed from the group.")
    discussion = ("This control addresses the establishment of system accounts, including group accounts (i.e., accounts shared by multiple individuals). Users may include, for example, organizational employees "
                  "or individuals to whom the organization has granted access rights. Organization-defined conditions for group membership may include, for example, criteria for membership based on roles and "
                  "attributes. Conditions for group membership may require approval by authorized individuals or roles. Information system account types can include, for example, individual, shared, group, "
                  "system, guest/anonymous, and temporary accounts.")
    related_controls = ("AC-1, AC-3, AC-4, AC-5, AC-6")

    actions = [
        "Review and update account management policies.",
        "Assign account managers to oversee account activities.",
        "Conduct regular audits of active accounts."
    ]
    
    milestones = [
        "Policy review completed - Estimated completion: 2024-09-01",
        "Account managers assigned - Estimated completion: 2024-09-15",
        "First audit conducted - Estimated completion: 2024-10-01"
    ]
    
    resources = [
        "Dedicated account management software.",
        "Training sessions for account managers.",
        "Additional personnel for audit processes."
    ]

    system_name = "System Name"
    owner = "Owner Name"

    if len(sys.argv) > 1:
        control_id = sys.argv[1]
        # Check for second command line argument for system name
        if len(sys.argv) > 2:
            system_name = sys.argv[2]
        # Check for third command line argument for owner
        if len(sys.argv) > 3:
            owner = sys.argv[3]
        # Read the CSV file
        df = pd.read_csv('controls.csv')
        control_row = df[df.iloc[:, 0] == control_id]

        if control_row.empty:
            print(f"Control ID {control_id} not found in CSV file.")
            return

        control = control_row.iloc[0, 1]
        control_text = control_row.iloc[0, 2]
        discussion = control_row.iloc[0, 3]
        related_controls = control_row.iloc[0, 4]
        if pd.isna(related_controls):
            related_controls = ""
        actions = ["Placeholder action 1", "Placeholder action 2"]
        milestones = ["Placeholder milestone 1 - Estimated completion: Date1", "Placeholder milestone 2 - Estimated completion: Date2"]
        resources = ["Placeholder resource 1", "Placeholder resource 2"]
    
    generate_document(control_id, control, control_text, discussion, related_controls, actions, milestones, resources, system_name, owner)

if __name__ == "__main__":
    main()
